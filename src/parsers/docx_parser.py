# DOCX Parser Module
"""
Word文档解析器
"""

from typing import Dict, Any
from pathlib import Path

from ..core.exceptions import DocumentParseError


class DocxParser:
    """Word文档解析器"""

    def __init__(self):
        """初始化Word解析器"""
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖库是否可用"""
        try:
            from docx import Document
            self._has_docx = True
        except ImportError:
            self._has_docx = False
            raise ImportError(
                "Word解析需要安装 python-docx。"
                "请运行: pip install python-docx"
            )

    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        解析Word文件

        Args:
            file_path: Word文件路径
            **kwargs: 额外参数

        Returns:
            Dict: 包含文本内容和元数据的字典
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentParseError(f"文件不存在", file_path)

        if path.suffix.lower() not in [".docx", ".doc"]:
            raise DocumentParseError(f"不是Word文件", file_path)

        # .doc格式需要额外处理，这里只支持.docx
        if path.suffix.lower() == ".doc":
            raise DocumentParseError(
                "暂不支持.doc格式，请将文件转换为.docx格式后上传",
                file_path
            )

        return self._parse_docx(file_path)

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析.docx文件"""
        from docx import Document

        doc = Document(file_path)

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格文本
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_text))
            if table_rows:
                tables_text.append("\n".join(table_rows))

        # 合并文本
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n【表格内容】\n" + "\n\n".join(tables_text)

        return {
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "char_count": len(full_text),
            "parser": "python-docx"
        }

    def parse_from_bytes(self, file_bytes: bytes, **kwargs) -> Dict[str, Any]:
        """
        从字节数据解析Word文件

        Args:
            file_bytes: Word文件的字节数据
            **kwargs: 额外参数

        Returns:
            Dict: 包含文本内容和元数据的字典
        """
        import io
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格文本
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_text))
            if table_rows:
                tables_text.append("\n".join(table_rows))

        # 合并文本
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n【表格内容】\n" + "\n\n".join(tables_text)

        return {
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "char_count": len(full_text),
            "parser": "python-docx"
        }
